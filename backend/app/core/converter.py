from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
import tempfile
import os
import platform
import matplotlib.font_manager as fm


class DocumentConverter:
    """문서 형식 변환을 담당하는 클래스"""
    
    def __init__(self):
        """한글 폰트 설정 초기화"""
        self.korean_font = self._get_korean_font()
        if self.korean_font:
            self._register_korean_font()
    
    def _get_korean_font(self):
        """시스템에서 사용 가능한 한글 폰트를 찾습니다."""
        try:
            # 우선순위 폰트 목록
            preferred_fonts = ['NanumGothic', 'Malgun Gothic', 'AppleGothic', 'Noto Sans CJK KR']
            
            available_fonts = {f.name: f.fname for f in fm.fontManager.ttflist}
            
            for font_name in preferred_fonts:
                if font_name in available_fonts:
                    return {'name': font_name, 'path': available_fonts[font_name]}
            
            # 한글 폰트 검색
            korean_fonts = {name: path for name, path in available_fonts.items() 
                          if any(keyword in name.lower() for keyword in ['gothic', 'nanum', 'malgun'])}
            
            if korean_fonts:
                font_name = list(korean_fonts.keys())[0]
                return {'name': font_name, 'path': korean_fonts[font_name]}
                
            return None
        except Exception:
            return None
    
    def _register_korean_font(self):
        """reportlab에 한글 폰트 등록"""
        try:
            if self.korean_font and os.path.exists(self.korean_font['path']):
                pdfmetrics.registerFont(TTFont('KoreanFont', self.korean_font['path']))
                return True
        except Exception:
            pass
        return False
    
    def to_docx(self, text: str) -> bytes:
        """텍스트를 DOCX 형식으로 변환합니다."""
        try:
            doc = Document()
            
            # 제목 추가
            title = doc.add_heading('PDF 요약', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 한글 폰트 설정
            if self.korean_font:
                title.runs[0].font.name = self.korean_font['name']
                title.runs[0]._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', self.korean_font['name'])
            
            # 텍스트를 문단별로 나누어 추가
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    # 한글 폰트 적용
                    if self.korean_font:
                        for run in p.runs:
                            run.font.name = self.korean_font['name']
                            run.font.size = Pt(11)
                            run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', self.korean_font['name'])
                        
                        # 기본 run이 없는 경우를 위한 처리
                        if not p.runs:
                            run = p.runs[0] if p.runs else p.add_run()
                            run.font.name = self.korean_font['name']
                            run.font.size = Pt(11)
                            run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', self.korean_font['name'])
            
            # BytesIO에 저장
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise ValueError(f"DOCX 변환 실패: {str(e)}")
    
    def to_pdf(self, text: str) -> bytes:
        """텍스트를 PDF 형식으로 변환합니다."""
        try:
            buffer = BytesIO()
            
            # PDF 문서 생성
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # 스타일 설정
            styles = getSampleStyleSheet()
            
            # 한글 폰트 설정
            font_name = 'KoreanFont' if self.korean_font else 'Helvetica'
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_name,
                fontSize=18,
                spaceAfter=30,
                alignment=1  # 중앙 정렬
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=12,
                spaceAfter=12,
                alignment=0,  # 왼쪽 정렬
                leftIndent=0,
                rightIndent=0,
                leading=16  # 줄 간격
            )
            
            # 컨텐츠 생성
            story = []
            
            # 제목 추가
            story.append(Paragraph("PDF 요약", title_style))
            story.append(Spacer(1, 12))
            
            # 본문 추가
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # HTML 특수문자 이스케이프
                    clean_paragraph = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(clean_paragraph, normal_style))
                    story.append(Spacer(1, 6))
            
            # PDF 생성
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise ValueError(f"PDF 변환 실패: {str(e)}")
    
    def to_txt(self, text: str) -> bytes:
        """텍스트를 TXT 형식으로 변환합니다."""
        try:
            # UTF-8 인코딩으로 텍스트 변환
            header = "=== PDF 요약 ===\n\n"
            content = header + text
            return content.encode('utf-8')
            
        except Exception as e:
            raise ValueError(f"TXT 변환 실패: {str(e)}")
    
    def to_html(self, text: str) -> bytes:
        """텍스트를 HTML 형식으로 변환합니다. (추가 기능)"""
        try:
            # 문단 분리
            paragraphs = text.split('\n\n')
            paragraph_html = ''.join(f'<p>{paragraph.strip()}</p>' for paragraph in paragraphs if paragraph.strip())
            
            html_template = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>PDF 요약</title>
    <style>
        body {{
            font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}
        p {{
            margin-bottom: 15px;
            text-align: justify;
        }}
        .container {{
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            padding: 30px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 PDF 요약</h1>
        {paragraph_html}
    </div>
</body>
</html>"""
            return html_template.encode('utf-8')
            
        except Exception as e:
            raise ValueError(f"HTML 변환 실패: {str(e)}") 